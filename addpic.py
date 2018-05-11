from PIL import Image
from docx import Document
from docx.shared import Cm

document = Document()

p = document.add_paragraph()
r = p.add_run()
r.add_text('Good Morning every body,This is my ')
picPath = 'D:/Development/Python/aa.png'
r.add_picture(picPath)
r.add_text(' do you like it?')

document.save('demo.docx')

'''
Here's my solution. It has the advantage on the first proposition that it surrounds 
the picture with a title (with style 'Header 1') and a section for additional comments. 
Note that you have to do the insertions in the reverse order they appear in the Word 
document.

This snippet is particularly useful if you want to programmatically insert pictures 
in an existing document.
'''

document = Document()

p = document.add_paragraph('Picture bullet section', 'List Bullet')
p = p.insert_paragraph_before('')
r = p.add_run()
r.add_picture(picPath)
p = p.insert_paragraph_before('My picture title', 'Heading 1')

paragraph = p.insert_paragraph_before('A bullet section', 'List Bullet')
paragraph = paragraph.insert_paragraph_before('')
paragraphRun = paragraph.add_run()

IMG_MAX_WIDTH = 17.5  # anciennement 19.5
SCREEN_DPI = 144  # on my 1920 x 1080' monitor

from win32api import GetSystemMetrics
import math

scnPxWidth = GetSystemMetrics(0)
print("Width ={}".format(scnPxWidth))
scnPxHeight = GetSystemMetrics(1)
print("Height ={}".format(scnPxHeight))
scnDPI = math.sqrt(int(scnPxWidth) ** 2 + int(scnPxHeight) ** 2) / 13.3
print("DPI={}".format(scnDPI))


picPath = 'D:/Development/Python/full_6_scn.jpg'


im = Image.open(picPath)
imgWidthPixel, height = im.size
imgWidthCm = imgWidthPixel / scnDPI * 2.54
paragraphRun.add_picture(picPath, width=Cm(min(IMG_MAX_WIDTH, imgWidthCm)))

paragraph = paragraph.insert_paragraph_before('A title', 'Heading 1')

document.save('demo_better.docx')
