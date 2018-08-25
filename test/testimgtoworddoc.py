import inspect
import os
import sys
import shutil
import unittest

currentdir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
parentdir = os.path.dirname(currentdir)
sys.path.insert(0, parentdir)
sys.path.insert(0,currentdir) # this instruction is necessary for successful importation of utilityfortest module when
                              # the test is executed standalone

import imgToWordDoc
from docx import Document

class TestImgToWordDoc(unittest.TestCase):
    def testSortNumberedStringsFunc(self):
        listOfFileNames = ['2.jpg', '1.png', '32.jpg', '8.png', 'name12.png']

        listOfFileNames.sort(key = imgToWordDoc.sortNumberedStringsFunc)
        self.assertEqual(['1.png', '2.jpg', '8.png', 'name12.png', '32.jpg'], listOfFileNames)


    def testSortNumberedStringsFuncInvalidFileName(self):
        listOfFileNames = ['2.jpg', '1.png', '32.jpg', '8.png', 'name.png']

        with self.assertRaises(NameError):
            listOfFileNames.sort(key = imgToWordDoc.sortNumberedStringsFunc)


    def testAddImagesAtEndOfDocumentWithImgInDirImgNumber(self):
        newWordDoc = Document()
        imgFileLst = ['1.jpg', 'name3.jpg', '4.jpg']
        insertedImgNumber = imgToWordDoc.addImagesAtEndOfDocument(newWordDoc, imgFileLst)
        targetWordDoc = 'newDocWhereImgWereAdded.docx'
        newWordDoc.save(targetWordDoc)
        self.assertEqual(3, insertedImgNumber)
        self.assertEqual(insertedImgNumber * 3, len(newWordDoc.paragraphs))

        self.assertEqual('1_title', newWordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', newWordDoc.paragraphs[2].text)

        self.assertEqual('name3_title', newWordDoc.paragraphs[3].text)
        self.assertEqual('name3_bullet', newWordDoc.paragraphs[5].text)

        self.assertEqual('4_title', newWordDoc.paragraphs[6].text)
        self.assertEqual('4_bullet', newWordDoc.paragraphs[8].text)

        os.remove(targetWordDoc)


    def testDetermineUniqueFileNameNoWordFileExist(self):
        wordFileName = "notExistFileName"
        wordFileNameWithExt = imgToWordDoc.determineUniqueFileName(wordFileName)
        self.assertEqual("notExistFileName.docx", wordFileNameWithExt)


    def testDetermineUniqueFileNameWordFileExistNoSuffixNumber(self):

        # create the situation where there's 'existingFileName.docx' only
        # in the current dir

        wordFileName = "existingFileName"

        fileNameExt = wordFileName + ".docx"

        with open(fileNameExt, 'w') as f:
            pass

        # perform the test

        wordFileNameWithExt = imgToWordDoc.determineUniqueFileName(wordFileName)
        self.assertEqual("existingFileName1.docx", wordFileNameWithExt)

        # clean up

        os.remove(fileNameExt)


    def testDetermineUniqueFileNameTwoWordFileExistOneWithSuffixNumber(self):

        # create the situation where there's 'existingFileName.docx' and 'existingFileName1.docx'
        # in the current dir

        wordFileName = "existingFileName"
        fileNameExt = wordFileName + ".docx"

        with open(fileNameExt, 'w') as f:
            pass

        wordFileNameSuffixOne = wordFileName + '1'
        fileNameExtSuffixOne = wordFileNameSuffixOne + ".docx"

        with open(fileNameExtSuffixOne, 'w') as f:
            pass

        # perform the test

        wordFileNameWithExt = imgToWordDoc.determineUniqueFileName(wordFileName)

        # clean up

        os.remove(fileNameExt)
        os.remove(fileNameExtSuffixOne)

        self.assertEqual("existingFileName2.docx", wordFileNameWithExt)

    def testGetFilesInTestDir(self):
        curDir = os.getcwd()

        filesInDir = imgToWordDoc.getFilesInDir(curDir)
        self.assertEqual(10, len(filesInDir))


    def testGetFilesInEmptyDir(self):
        os.makedirs('empty')
        os.chdir('empty')
        curDir = os.getcwd()

        filesInDir = imgToWordDoc.getFilesInDir(curDir)
        self.assertEqual(0, len(filesInDir))

        os.chdir('..')
        os.removedirs('empty')


    def testFilterAndSortImageFileNames(self):
        curDir = os.getcwd()
        imgFileLst = imgToWordDoc.filterAndSortImageFileNames(curDir)
        self.assertEqual(['1.jpg', 'name3.jpg', '4.jpg'],imgFileLst)


    def testFilterAndSortImageFileNamesWithImageNumberToAdd(self):
        curDir = os.getcwd()
        imageNumbersToAdd = [1, 3]
        imgFileLst = imgToWordDoc.filterAndSortImageFileNames(curDir, imageNumbersToAdd)
        self.assertEqual(['1.jpg', 'name3.jpg'],imgFileLst)


    def testFilterAndSortImageFileNamesWithInvalidFileName(self):
        '''
        Current dir contains image file whose name contains no number.
        :return:
        '''
        curDir = os.getcwd()

        invalidFileName = 'invalFileName.jpg'

        with open(invalidFileName, 'w') as f:
            pass

        with self.assertRaises(NameError):
            imgToWordDoc.filterAndSortImageFileNames(curDir)

        os.remove(invalidFileName)


    def testDetermineInsertionPointInExistingWordDocWithOneParagraph(self):
        wordDoc = Document('oneImg.docx')
        insertionPoint = 1
        firstParagraph = wordDoc.paragraphs[0]

        titleString = firstParagraph.text
        self.assertEqual('My picture title', titleString)
        self.assertEqual(titleString, imgToWordDoc.determineInsertionPoint(insertionPoint, wordDoc).text)


    def testDetermineInsertionPointInExistingWordDocWithTwoParagraphsPos1(self):
        wordDoc = Document('twoImg.docx')
        insertionPoint = 1
        firstParagraph = wordDoc.paragraphs[0]

        titleString = firstParagraph.text
        self.assertEqual('My picture one title', titleString)
        self.assertEqual(titleString, imgToWordDoc.determineInsertionPoint(insertionPoint, wordDoc).text)


    def testDetermineInsertionPointInExistingWordDocWithTwoParagraphsPos2(self):
        wordDoc = Document('twoImg.docx')
        insertionPoint = 2
        firstParagraph = wordDoc.paragraphs[3] # 2 x (header + image + bullet points section

        titleString = firstParagraph.text
        self.assertEqual('My picture two title', titleString)
        self.assertEqual(titleString, imgToWordDoc.determineInsertionPoint(insertionPoint, wordDoc).text)


    def testDetermineInsertionPointInExistingWordDocWithTwoParagraphsPos0(self):
        '''
        Insertion point 0 means the insertrion occurs at the end of document.
        :return:
        '''
        wordDoc = Document('twoImg.docx')
        insertionPoint = 0
        self.assertIsNone(imgToWordDoc.determineInsertionPoint(insertionPoint, wordDoc))


    def testDetermineInsertionPointInExistingWordDocWithTwoParagraphsPos3(self):
        '''
        Insertion point 3 exceeds the number of 'Heading1' paragraphs of the document and
        will cause the insertrion to occur at the end of document.
        :return:
        '''
        wordDoc = Document('twoImg.docx')
        insertionPoint = 3
        self.assertIsNone(imgToWordDoc.determineInsertionPoint(insertionPoint, wordDoc))


    def testInsertImagesBeforeParagraphTwoInTwoParagraphsDoc(self):
        wordDoc = Document('twoImgForInsertion.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        secondHeaderParagraph = wordDoc.paragraphs[3] # 2 x (header + image + bullet points section
        imgFileLst = ['1.jpg', 'name3.jpg', '4.jpg']
        insertedImgNumber = imgToWordDoc.insertImagesBeforeParagraph(secondHeaderParagraph, imgFileLst)
        targetWordDoc = 'twoImgForInsertion1.docx'
        wordDoc.save(targetWordDoc)
        self.assertEqual(3, insertedImgNumber)
        self.assertEqual(initialParagraphNumber + 3 * 3, len(wordDoc.paragraphs))

        self.assertEqual('My picture one title', wordDoc.paragraphs[0].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[2].text)

        self.assertEqual('1_title', wordDoc.paragraphs[3].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[6].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('4_title', wordDoc.paragraphs[9].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[12].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[14].text)

        os.remove(targetWordDoc)


    def testInsertImagesBeforeParagraphOneInTwoParagraphsDoc(self):
        wordDoc = Document('twoImgForInsertion.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        firstHeaderParagraph = wordDoc.paragraphs[0] # 2 x (header + image + bullet points section
        imgFileLst = ['1.jpg', 'name3.jpg', '4.jpg']
        insertedImgNumber = imgToWordDoc.insertImagesBeforeParagraph(firstHeaderParagraph, imgFileLst)
        targetWordDoc = 'twoImgForInsertion1.docx'
        wordDoc.save(targetWordDoc)
        self.assertEqual(3, insertedImgNumber)
        self.assertEqual(initialParagraphNumber + 3 * 3, len(wordDoc.paragraphs))

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[3].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('4_title', wordDoc.paragraphs[6].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('My picture one title', wordDoc.paragraphs[9].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[11].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[12].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[14].text)

        os.remove(targetWordDoc)


    def testInsertImagesBeforeParagraphOneInOneParagraphsDoc(self):
        wordDoc = Document('oneImgForInsertion.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        firstHeaderParagraph = wordDoc.paragraphs[0] # 2 x (header + image + bullet points section
        imgFileLst = ['1.jpg', 'name3.jpg', '4.jpg']
        insertedImgNumber = imgToWordDoc.insertImagesBeforeParagraph(firstHeaderParagraph, imgFileLst)
        targetWordDoc = 'oneImgForInsertion1.docx'
        wordDoc.save(targetWordDoc)
        self.assertEqual(3, insertedImgNumber)
        self.assertEqual(initialParagraphNumber + 3 * 3, len(wordDoc.paragraphs))

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[3].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('4_title', wordDoc.paragraphs[6].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('My picture one title', wordDoc.paragraphs[9].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[11].text)

        os.remove(targetWordDoc)


    def testInsertImagesBeforeParagraphOneInEmptiedDoc(self):
        '''
        Inserting into a Word document which were emptied works since wordDoc.paragraphs[0]
        returns a paragraph with style 'Normal'
        '''
        wordDoc = Document('emptyDocForInsertion.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        firstHeaderParagraph = wordDoc.paragraphs[0] # 2 x (header + image + bullet points section
        imgFileLst = ['1.jpg', 'name3.jpg', '4.jpg']
        insertedImgNumber = imgToWordDoc.insertImagesBeforeParagraph(firstHeaderParagraph, imgFileLst)
        targetWordDoc = 'emptyDocForInsertion1.docx'
        wordDoc.save(targetWordDoc)
        self.assertEqual(3, insertedImgNumber)
        self.assertEqual(initialParagraphNumber + 3 * 3, len(wordDoc.paragraphs))

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[3].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('4_title', wordDoc.paragraphs[6].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[8].text)

        os.remove(targetWordDoc)


    def testInsertImagesBeforeParagraphInNewDoc(self):
        '''
        Inserting into an new Word document is not possible because the document has no
        paragraph.
        '''
        wordDoc = Document()
        initialParagraphNumber = len(wordDoc.paragraphs)
        self.assertEqual(0, initialParagraphNumber)


    def testCreateOrUpdateWordDocWithImgInDirIncrementFileNameCreationMode(self):
        imgToWordDoc.createOrUpdateWordDocWithImgInDir()
        imgToWordDoc.createOrUpdateWordDocWithImgInDir()

        filesInDirList = os.listdir(currentdir)

        wordDoc = Document('test.docx')
        wordDoc1 = Document('test1.docx')

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(currentdir + '\\test.docx')
        os.remove(currentdir + '\\test1.docx')

        self.assertIn('test.docx', filesInDirList)
        self.assertIn('test1.docx', filesInDirList)

        self.assertEqual(3, len(wordDoc.paragraphs) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[3].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('4_title', wordDoc.paragraphs[6].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual(3, len(wordDoc1.paragraphs) / 3)

        self.assertEqual('1_title', wordDoc1.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc1.paragraphs[2].text)

        self.assertEqual('name3_title', wordDoc1.paragraphs[3].text)
        self.assertEqual('name3_bullet', wordDoc1.paragraphs[5].text)

        self.assertEqual('4_title', wordDoc1.paragraphs[6].text)
        self.assertEqual('4_bullet', wordDoc1.paragraphs[8].text)


    def testCreateOrUpdateWordDocWithImgInDirImgCreationMode(self):
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir()
        wordFilePathName = currentdir + '\\test.docx'
        wordDoc = Document(wordFilePathName)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(wordFilePathName)

        self.assertEqual("test.docx file created with 3 image(s). Manually add auto numbering to the 'Header 1' / 'Titre 1' style !", returnedInfo)
        self.assertEqual(9, len(wordDoc.paragraphs)) # 3 headers + 3 images + 3 bullet points sections

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[3].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('4_title', wordDoc.paragraphs[6].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[8].text)


    def testCreateOrUpdateWordDocWithImgInDirImgCreationModeDocNameSpecified(self):
        docName = 'monDocTest'
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(docName)])
        wordFilePathName = currentdir + '\\{}.docx'.format(docName)
        wordDoc = Document(wordFilePathName)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(wordFilePathName)

        self.assertEqual(docName + ".docx file created with 3 image(s). Manually add auto numbering to the 'Header 1' / 'Titre 1' style !", returnedInfo)
        self.assertEqual(9, len(wordDoc.paragraphs)) # 3 headers + 3 images + 3 bullet points sections

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[3].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('4_title', wordDoc.paragraphs[6].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[8].text)


    def testCreateOrUpdateWordDocInsertImagesBeforeParagraphTwoInTwoParagraphsDoc(self):
        initialWordDocNameNoExt = 'twoImgForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i2'])
        finalWordDoc = 'twoImgForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)

        self.assertEqual("Inserted 3 image(s) before header 2 in document twoImgForInsertion.docx and saved the result to twoImgForInsertion1.docx.", returnedInfo)
        self.assertEqual(3, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('My picture one title', wordDoc.paragraphs[0].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[2].text)

        self.assertEqual('1_title', wordDoc.paragraphs[3].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[6].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('4_title', wordDoc.paragraphs[9].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[12].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[14].text)


    def testCreateOrUpdateWordDocInsertImagesBeforeParagraphOneInTwoParagraphsDoc(self):
        initialWordDocNameNoExt = 'twoImgForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i1'])
        finalWordDoc = 'twoImgForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)

        self.assertEqual("Inserted 3 image(s) before header 1 in document twoImgForInsertion.docx and saved the result to twoImgForInsertion1.docx.", returnedInfo)
        self.assertEqual(3, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[3].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('4_title', wordDoc.paragraphs[6].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('My picture one title', wordDoc.paragraphs[9].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[11].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[12].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[14].text)


    def testCreateOrUpdateWordDocInsertImagesBeforeParagraphOneInOneParagraphsDoc(self):
        initialWordDocNameNoExt = 'oneImgForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i1'])
        finalWordDoc = 'oneImgForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)

        self.assertEqual("Inserted 3 image(s) before header 1 in document oneImgForInsertion.docx and saved the result to oneImgForInsertion1.docx.", returnedInfo)
        self.assertEqual(3, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[3].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('4_title', wordDoc.paragraphs[6].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('My picture one title', wordDoc.paragraphs[9].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[11].text)


    def testCreateOrUpdateWordDocInsertImagesBeforeParagraphOneInEmptiedDoc(self):
        '''
        Inserting into a Word document which were emptied works since wordDoc.paragraphs[0]
        returns a paragraph with style 'Normal'
        '''
        initialWordDocNameNoExt = 'emptyDocForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i1'])
        finalWordDoc = 'emptyDocForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)

        self.assertEqual("Added 3 image(s) at end of document emptyDocForInsertion.docx and saved the result to emptyDocForInsertion1.docx. Although insertion position 1 was provided, no header paragraph was available at this position and the images were added at the end of the document !", returnedInfo)
        self.assertEqual(3, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[1].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[3].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[4].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[6].text)

        self.assertEqual('4_title', wordDoc.paragraphs[7].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[9].text)


    def testCreateOrUpdateWordDocInsertImagesAtEndInTwoParagraphsDoc(self):
        initialWordDocNameNoExt = 'twoImgForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i0'])
        finalWordDoc = 'twoImgForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)

        self.assertEqual("Added 3 image(s) at end of document twoImgForInsertion.docx and saved the result to twoImgForInsertion1.docx.", returnedInfo)
        self.assertEqual(3, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[6].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[9].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('4_title', wordDoc.paragraphs[12].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[14].text)

        self.assertEqual('My picture one title', wordDoc.paragraphs[0].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[2].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[3].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[5].text)


    def testCreateOrUpdateWordDocInsertImagesAtPosExceedingHeaderNumberInTwoParagraphsDoc(self):
        initialWordDocNameNoExt = 'twoImgForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i20'])
        finalWordDoc = 'twoImgForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)

        self.assertEqual("Added 3 image(s) at end of document twoImgForInsertion.docx and saved the result to twoImgForInsertion1.docx. Although insertion position 20 was provided, no header paragraph was available at this position and the images were added at the end of the document !", returnedInfo)
        self.assertEqual(3, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[6].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[9].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('4_title', wordDoc.paragraphs[12].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[14].text)

        self.assertEqual('My picture one title', wordDoc.paragraphs[0].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[2].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[3].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[5].text)


    def testCreateOrUpdateWordDocInsertImagesAtEndOfEmptiedDoc(self):
        '''
        Inserting into a Word document which were emptied works since wordDoc.paragraphs[0]
        returns a paragraph with style 'Normal'
        '''
        initialWordDocNameNoExt = 'emptyDocForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i0'])
        finalWordDoc = 'emptyDocForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)

        self.assertEqual("Added 3 image(s) at end of document emptyDocForInsertion.docx and saved the result to emptyDocForInsertion1.docx.", returnedInfo)

        self.assertEqual(3, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[1].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[3].text)

        self.assertEqual('name3_title', wordDoc.paragraphs[4].text)
        self.assertEqual('name3_bullet', wordDoc.paragraphs[6].text)

        self.assertEqual('4_title', wordDoc.paragraphs[7].text)
        self.assertEqual('4_bullet', wordDoc.paragraphs[9].text)


    def testExplodeImageNumbersListSimple(self):
        imageNumberSpecs = ['1', '3', '4', '2', '3', '4']
        self.assertEqual([1, 2, 3, 4], imgToWordDoc.explodeImageNumbersList(imageNumberSpecs))


    def testExplodeImageNumbersListSimpleAsString(self):
        imageNumberSpecs = ['1 3 4 2 3 4']
        self.assertEqual([1, 2, 3, 4], imgToWordDoc.explodeImageNumbersList(imageNumberSpecs))


    def testExplodeImageNumbersListAsString(self):
        imageNumberSpecs = ['1 3 4 2 13-14']
        self.assertEqual([1, 2, 3, 4, 13, 14], imgToWordDoc.explodeImageNumbersList(imageNumberSpecs))


    def testExplodeImageNumbersListOneElement(self):
        imageNumberSpecs = ['1']
        self.assertEqual([1], imgToWordDoc.explodeImageNumbersList(imageNumberSpecs))


    def testExplodeNumberSpec(self):
        imageNumberSpec = '1-3'
        self.assertEqual([1, 2, 3], imgToWordDoc.explodeNumberSpec(imageNumberSpec))


    def testExplodeNumberSpecOneToOne(self):
        imageNumberSpec = '1-1'
        self.assertEqual([1], imgToWordDoc.explodeNumberSpec(imageNumberSpec))


    def testExplodeNumberSpecRev(self):
        imageNumberSpec = '3-1'
        self.assertEqual([1, 2, 3], imgToWordDoc.explodeNumberSpec(imageNumberSpec))


    def testExplodeNumberSpecRev(self):
        imageNumberSpec = '9-12'
        self.assertEqual([9, 10, 11, 12], imgToWordDoc.explodeNumberSpec(imageNumberSpec))


    def testExplodeImageNumbersList(self):
        imageNumberSpecs = ['1', '3', '4', '2-7', '9-12']
        self.assertEqual([1, 2, 3, 4, 5, 6, 7, 9, 10, 11, 12], imgToWordDoc.explodeImageNumbersList(imageNumberSpecs))


    def testFilterAccordingToNumber(self):
        unfilteredImgFileNameLst = ['1.png', 'name3.jpg', '2nom.jpg', '4.jpg']
        numberLst = [1, 2, 3, 5]
        imgFileLst = imgToWordDoc.filterAccordingToNumber(unfilteredImgFileNameLst, numberLst)
        self.assertEqual(['1.png', 'name3.jpg', '2nom.jpg'],imgFileLst)


    def testFilterAccordingToNumberInvalidImgFileName(self):
        unfilteredImgFileNameLst = ['1.png', 'name3.jpg', '2nom.jpg', 'invalid_no_number.jpg', '4.jpg']
        numberLst = [1, 2, 3, 5]
        with self.assertRaises(NameError):
            imgToWordDoc.filterAccordingToNumber(unfilteredImgFileNameLst, numberLst)


    def testFilterAccordingToNumberNoMatch(self):
        unfilteredImgFileNameLst = ['1.png', 'name3.jpg', '2nom.jpg', '4.jpg']
        numberLst = [5, 6]
        imgFileLst = imgToWordDoc.filterAccordingToNumber(unfilteredImgFileNameLst, numberLst)
        self.assertEqual([],imgFileLst)


    def testFilterAccordingToNumberEmptyFileNameLst(self):
        unfilteredImgFileNameLst = []
        numberLst = [1, 2, 5]
        imgFileLst = imgToWordDoc.filterAccordingToNumber(unfilteredImgFileNameLst, numberLst)
        self.assertEqual([],imgFileLst)


    def testFilterAccordingToNumberEmptyNumberLst(self):
        unfilteredImgFileNameLst = ['1.png', 'name3.jpg', '2nom.jpg', '4.jpg']
        numberLst = []
        imgFileLst = imgToWordDoc.filterAccordingToNumber(unfilteredImgFileNameLst, numberLst)
        self.assertEqual([],imgFileLst)


    def testCreateOrUpdateWordDocInsertSelectedImagesListWithSpacesAtEndInTwoParagraphsDoc(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'twoImgForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i0', '-p 1 2 5-7 10-9 12'])
        finalWordDoc = 'twoImgForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual("Added 6 image(s) at end of document twoImgForInsertion.docx and saved the result to twoImgForInsertion1.docx.", returnedInfo)

        self.assertEqual(6, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[6].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('5_title', wordDoc.paragraphs[9].text)
        self.assertEqual('5_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('6_title', wordDoc.paragraphs[12].text)
        self.assertEqual('6_bullet', wordDoc.paragraphs[14].text)

        self.assertEqual('7_title', wordDoc.paragraphs[15].text)
        self.assertEqual('7_bullet', wordDoc.paragraphs[17].text)

        self.assertEqual('9_title', wordDoc.paragraphs[18].text)
        self.assertEqual('9_bullet', wordDoc.paragraphs[20].text)

        self.assertEqual('10twoDigit_title', wordDoc.paragraphs[21].text)
        self.assertEqual('10twoDigit_bullet', wordDoc.paragraphs[23].text)

        self.assertEqual('My picture one title', wordDoc.paragraphs[0].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[2].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[3].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[5].text)


    def testCreateOrUpdateWordDocInsertSelectedImagesListWithCommasAtEndInTwoParagraphsDoc(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'twoImgForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i0', '-p 1, 2 ,5-7,10-9'])
        finalWordDoc = 'twoImgForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual(
            "Added 6 image(s) at end of document twoImgForInsertion.docx and saved the result to twoImgForInsertion1.docx.",
            returnedInfo)

        self.assertEqual(6, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[6].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('5_title', wordDoc.paragraphs[9].text)
        self.assertEqual('5_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('6_title', wordDoc.paragraphs[12].text)
        self.assertEqual('6_bullet', wordDoc.paragraphs[14].text)

        self.assertEqual('7_title', wordDoc.paragraphs[15].text)
        self.assertEqual('7_bullet', wordDoc.paragraphs[17].text)

        self.assertEqual('9_title', wordDoc.paragraphs[18].text)
        self.assertEqual('9_bullet', wordDoc.paragraphs[20].text)

        self.assertEqual('10twoDigit_title', wordDoc.paragraphs[21].text)
        self.assertEqual('10twoDigit_bullet', wordDoc.paragraphs[23].text)

        self.assertEqual('My picture one title', wordDoc.paragraphs[0].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[2].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[3].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[5].text)


    def deleteFiles(self, fileNamesList):
        for fileName in fileNamesList:
            os.remove(fileName)


    def copyDirContent(self, fromDir, toDir):
        '''
        Copy the fromDir files to toDir and returns the list of copied files.

        :param fromDir:
        :param toDir:

        :return: list of copied file names without path
        '''
        fileNameList = os.listdir(fromDir)

        for fileName in fileNameList:
            filePathName = os.path.join(fromDir, fileName)
            if (os.path.isfile(filePathName)):
                shutil.copy(filePathName, toDir)

        return fileNameList


    def testCreateOrUpdateWordDocInsertSelectedImagesListWithSpacesAtStartInTwoParagraphsDoc(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'twoImgForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i1', '-p 1 2 5-7 10-9 12'])
        finalWordDoc = 'twoImgForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual(
            "Inserted 6 image(s) before header 1 in document twoImgForInsertion.docx and saved the result to twoImgForInsertion1.docx.",
            returnedInfo)

        self.assertEqual(6, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('5_title', wordDoc.paragraphs[3].text)
        self.assertEqual('5_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('6_title', wordDoc.paragraphs[6].text)
        self.assertEqual('6_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('7_title', wordDoc.paragraphs[9].text)
        self.assertEqual('7_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('9_title', wordDoc.paragraphs[12].text)
        self.assertEqual('9_bullet', wordDoc.paragraphs[14].text)

        self.assertEqual('10twoDigit_title', wordDoc.paragraphs[15].text)
        self.assertEqual('10twoDigit_bullet', wordDoc.paragraphs[17].text)

        self.assertEqual('My picture one title', wordDoc.paragraphs[18].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[20].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[21].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[23].text)


    def testCreateOrUpdateWordDocInsertSelectedImagesListWithSpacesBeforeHeaderTwoInTwoParagraphsDoc(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'twoImgForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i2', '-p 1 2 5-7 10-9 12'])
        finalWordDoc = 'twoImgForInsertion1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual(
            "Inserted 6 image(s) before header 2 in document twoImgForInsertion.docx and saved the result to twoImgForInsertion1.docx.",
            returnedInfo)

        self.assertEqual(6, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('My picture one title', wordDoc.paragraphs[0].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[2].text)

        self.assertEqual('1_title', wordDoc.paragraphs[3].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('5_title', wordDoc.paragraphs[6].text)
        self.assertEqual('5_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('6_title', wordDoc.paragraphs[9].text)
        self.assertEqual('6_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('7_title', wordDoc.paragraphs[12].text)
        self.assertEqual('7_bullet', wordDoc.paragraphs[14].text)

        self.assertEqual('9_title', wordDoc.paragraphs[15].text)
        self.assertEqual('9_bullet', wordDoc.paragraphs[17].text)

        self.assertEqual('10twoDigit_title', wordDoc.paragraphs[18].text)
        self.assertEqual('10twoDigit_bullet', wordDoc.paragraphs[20].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[21].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[23].text)


    def testCreateOrUpdateWordDocAddSelectedImagesListWithSpacesInNewDocumentWithInsertionPoint(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'newDocForAddingSelectedImages'
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i1', '-p 1 2 5-7 10-9 12'])
        finalWordDoc = initialWordDocNameNoExt + '.docx'
        wordDoc = Document(finalWordDoc)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual(
            "Added 6 image(s) at end of document newDocForAddingSelectedImages.docx and saved the result to newDocForAddingSelectedImages.docx. Although insertion position 1 was provided, no header paragraph was available at this position and the images were added at the end of the document !",
            returnedInfo)
        finalParagraphNumber = len(wordDoc.paragraphs)

        self.assertEqual(6, finalParagraphNumber / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('5_title', wordDoc.paragraphs[3].text)
        self.assertEqual('5_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('6_title', wordDoc.paragraphs[6].text)
        self.assertEqual('6_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('7_title', wordDoc.paragraphs[9].text)
        self.assertEqual('7_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('9_title', wordDoc.paragraphs[12].text)
        self.assertEqual('9_bullet', wordDoc.paragraphs[14].text)

        self.assertEqual('10twoDigit_title', wordDoc.paragraphs[15].text)
        self.assertEqual('10twoDigit_bullet', wordDoc.paragraphs[17].text)


    def testCreateOrUpdateWordDocAddSelectedImagesListWithSpacesInNewDocumentNoDocNameNoInsertionPoint(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'test' #name of containing dir since parm -d not specified
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(['-p 1 2 5-7 10-9 12'])
        finalWordDoc = initialWordDocNameNoExt + '.docx'
        wordDoc = Document(finalWordDoc)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual(
            "test.docx file created with 6 image(s). Manually add auto numbering to the 'Header 1' / 'Titre 1' style !",
            returnedInfo)
        finalParagraphNumber = len(wordDoc.paragraphs)

        self.assertEqual(6, finalParagraphNumber / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('5_title', wordDoc.paragraphs[3].text)
        self.assertEqual('5_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('6_title', wordDoc.paragraphs[6].text)
        self.assertEqual('6_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('7_title', wordDoc.paragraphs[9].text)
        self.assertEqual('7_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('9_title', wordDoc.paragraphs[12].text)
        self.assertEqual('9_bullet', wordDoc.paragraphs[14].text)

        self.assertEqual('10twoDigit_title', wordDoc.paragraphs[15].text)
        self.assertEqual('10twoDigit_bullet', wordDoc.paragraphs[17].text)


    def testCreateOrUpdateWordDocAddSelectedImagesListWithSpacesInNewDocumentNoInsertionPoint(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'newDocForAddingSelectedImages'
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-p 1 2 5-7 10-9 12'])
        finalWordDoc = initialWordDocNameNoExt + '.docx'
        wordDoc = Document(finalWordDoc)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual(
            "newDocForAddingSelectedImages.docx file created with 6 image(s). Manually add auto numbering to the 'Header 1' / 'Titre 1' style !",
            returnedInfo)
        finalParagraphNumber = len(wordDoc.paragraphs)

        self.assertEqual(6, finalParagraphNumber / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[0].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[2].text)

        self.assertEqual('5_title', wordDoc.paragraphs[3].text)
        self.assertEqual('5_bullet', wordDoc.paragraphs[5].text)

        self.assertEqual('6_title', wordDoc.paragraphs[6].text)
        self.assertEqual('6_bullet', wordDoc.paragraphs[8].text)

        self.assertEqual('7_title', wordDoc.paragraphs[9].text)
        self.assertEqual('7_bullet', wordDoc.paragraphs[11].text)

        self.assertEqual('9_title', wordDoc.paragraphs[12].text)
        self.assertEqual('9_bullet', wordDoc.paragraphs[14].text)

        self.assertEqual('10twoDigit_title', wordDoc.paragraphs[15].text)
        self.assertEqual('10twoDigit_bullet', wordDoc.paragraphs[17].text)


    def testCreateOrUpdateWordDocAddSelectedImagesListWithSpacesInEmptiedDocument(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'emptyDocForInsertion'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i1', '-p 1 2 5-7 10-9 12'])
        finalWordDoc = initialWordDocNameNoExt + '1.docx'
        wordDoc = Document(finalWordDoc)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual(
            "Added 6 image(s) at end of document emptyDocForInsertion.docx and saved the result to emptyDocForInsertion1.docx. Although insertion position 1 was provided, no header paragraph was available at this position and the images were added at the end of the document !",
            returnedInfo)
        finalParagraphNumber = len(wordDoc.paragraphs)

        self.assertEqual(6, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('1_title', wordDoc.paragraphs[1].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[3].text)

        self.assertEqual('5_title', wordDoc.paragraphs[4].text)
        self.assertEqual('5_bullet', wordDoc.paragraphs[6].text)

        self.assertEqual('6_title', wordDoc.paragraphs[7].text)
        self.assertEqual('6_bullet', wordDoc.paragraphs[9].text)

        self.assertEqual('7_title', wordDoc.paragraphs[10].text)
        self.assertEqual('7_bullet', wordDoc.paragraphs[12].text)

        self.assertEqual('9_title', wordDoc.paragraphs[13].text)
        self.assertEqual('9_bullet', wordDoc.paragraphs[15].text)

        self.assertEqual('10twoDigit_title', wordDoc.paragraphs[16].text)
        self.assertEqual('10twoDigit_bullet', wordDoc.paragraphs[18].text)


    def testCreateOrUpdateWordDocAddSelectedImagesListWithInvalidImgFileName(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        invalidFileName = 'invalFileName.jpg'

        with open(invalidFileName, 'w') as f:
            pass

        initialWordDocNameNoExt = 'emptyDocForInsertion'
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i1', '-p 1 2 5-7 10-9 12'])

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        self.deleteFiles(copiedFileNamesList)
        os.remove(invalidFileName)

        self.assertEqual(returnedInfo, "ERROR - Invalid img file name encountered: invalFileName.jpg. Img file names must contain a number for them to be inserted in the right order (depends on this number) !")


    def testCreateOrUpdateWordDocAddSelectedImageInNewDocumentNoDocNameWithInsertionPoint(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'test' #name of containing dir since parm -d not specified
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(['-i 1', '-p 10'])
        finalWordDoc = initialWordDocNameNoExt + '.docx'

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual(
            'ERROR - When specifying an insertion position (-i), an existing document name (-d) must be provided !',
            returnedInfo)


    def testCreateOrUpdateWordDocInsertSelectedImagesListWithSpacesBeforeHeaderTwoInTwoParagraphsDocWhereANewParagraphWaaAddedManually(self):
        testImgDir = currentdir + "\\images"
        copiedFileNamesList = self.copyDirContent(testImgDir, currentdir)

        initialWordDocNameNoExt = 'twoImgForInsertionWithParagraphAddedManually'
        wordDoc = Document(initialWordDocNameNoExt + '.docx')
        initialParagraphNumber = len(wordDoc.paragraphs)
        returnedInfo = imgToWordDoc.createOrUpdateWordDocWithImgInDir(["-d{}".format(initialWordDocNameNoExt), '-i2', '-p 1 2 5-7 10-9 12'])
        finalWordDoc = 'twoImgForInsertionWithParagraphAddedManually1.docx'
        wordDoc = Document(finalWordDoc)
        finalParagraphNumber = len(wordDoc.paragraphs)

        # clean up files written on disc before assertion checking so that if an assertion fails,
        # this does not impact the other tests !
        os.remove(finalWordDoc)
        self.deleteFiles(copiedFileNamesList)

        self.assertEqual(
            "Inserted 6 image(s) before header 2 in document twoImgForInsertionWithParagraphAddedManually.docx and saved the result to twoImgForInsertionWithParagraphAddedManually1.docx.",
            returnedInfo)

        self.assertEqual(6, (finalParagraphNumber - initialParagraphNumber) / 3)

        self.assertEqual('My picture one title', wordDoc.paragraphs[2].text)
        self.assertEqual('Picture one bullet section', wordDoc.paragraphs[4].text)

        self.assertEqual('1_title', wordDoc.paragraphs[5].text)
        self.assertEqual('1_bullet', wordDoc.paragraphs[7].text)

        self.assertEqual('5_title', wordDoc.paragraphs[8].text)
        self.assertEqual('5_bullet', wordDoc.paragraphs[10].text)

        self.assertEqual('6_title', wordDoc.paragraphs[11].text)
        self.assertEqual('6_bullet', wordDoc.paragraphs[13].text)

        self.assertEqual('7_title', wordDoc.paragraphs[14].text)
        self.assertEqual('7_bullet', wordDoc.paragraphs[16].text)

        self.assertEqual('9_title', wordDoc.paragraphs[17].text)
        self.assertEqual('9_bullet', wordDoc.paragraphs[19].text)

        self.assertEqual('10twoDigit_title', wordDoc.paragraphs[20].text)
        self.assertEqual('10twoDigit_bullet', wordDoc.paragraphs[22].text)

        self.assertEqual('My picture two title', wordDoc.paragraphs[23].text)
        self.assertEqual('Picture two bullet section', wordDoc.paragraphs[25].text)
