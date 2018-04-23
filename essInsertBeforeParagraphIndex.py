import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    i = 1

    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            if i == 2:
                #adding 2 images to existing doc
                insertImageBefore(para, 'D:/Development/Python/aa.png')
                insertImageBefore(para, 'D:/Development/Python/ab.jpg')
            i += 1
    doc.save(filename)

    return i


def insertImageBefore(paragraph, imageFileName):
    paragraph.insert_paragraph_before('A', 'Heading 1')
    newParagraph = paragraph.insert_paragraph_before('')
    newParagraphRun = newParagraph.add_run()
    newParagraphRun.add_picture(imageFileName)
    paragraph.insert_paragraph_before('A', 'List Bullet')


print(getText('ImgToWordDoc12.docx'))

