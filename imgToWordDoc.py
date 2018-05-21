import os, sys
import re
from win32api import GetSystemMetrics
import math
import os.path as curDir
from PIL import Image
from docx import Document
from docx.shared import Cm
import argparse

IMG_MAX_WIDTH = 17.5  # anciennement 19.5
LATERAL_MARGIN = 2  # anciennement 1

# calculating SCREEN_DPI constant
scnPxWidth = GetSystemMetrics(0)
scnPxHeight = GetSystemMetrics(1)

SCREEN_DPI = int(math.sqrt(scnPxWidth ** 2 + scnPxHeight ** 2) / 13.3) # 144 old value set on my 1920 x 1080' monitor

WORD_FILE_EXT = ".docx"
HEADING_ONE_STYLE_NAME_ENGLISH = 'Heading1'
HEADING_ONE_STYLE_NAME_FRENCH = 'Titre1'


def getCommandLineArgs(argList):
    '''
    Uses argparse to acquire the user optional command line arguments.

    :param argList: were acquired from sys.argv or set by test code
    :return: document name (may be None), insertion point and image numbers list to add/insert
    '''
    parser = argparse.ArgumentParser(
        description="Adds or inserts all or part of the images contained in the current dir to a Word document. Each image " \
                    "is added in a new paragraph. To facilitate further edition, the image " \
                    "is preceded by a header line and followed by a bullet point section. " \
                    "The images are added according to the ascending order of the number contained in their " \
                    "file name. An error will occur if one of the image file name does not "
                    "contain a number (valid image file names are: 1.jpg, image2.jpg, 3.png, ...). " \
                    "If no document name is specified, the created document has " \
                    "the same name as the containing dir. An existing document with " \
                    "same name is never overwritten. Instead, a new document with a " \
                    "name incremented by 1 (i.e. myDoc1.docx, myDoc2.docx, ...) " \
                    "is created. " \
                    "Using the utility in add mode, i.e. without specifying an insertion " \
                    "point, creates a new document in which the specified images will be added. " \
                    "If the current dir already contains a document with images and comments you " \
                    "want to keep, use the insertion parameter which will insert the new images at " \
                    "the specified position and preserve the initial content. " \
                    "Without using the -p parameter, all the images of the current dir are collected " \
                    "for the addition/insertion. -p enables to specify precisely the images to select. ")
    parser.add_argument("-d", "--document", nargs="?", help="existing document to which the images are " \
                                                            "to be added. For your convenience, the initial document is " \
                                                            "not modified. Instead, the original document is copied with a " \
                                                            "name incremented by one and the images are added/inserted to the copy.")
    parser.add_argument("-i", "--insertionPos", type=int, nargs="?",
                        help="paragraph number BEFORE which to insert the " \
                             "images. 1 --> start of document (before paragraph 1). " \
                             "0 --> end of document. ")
    parser.add_argument("-p", "--pictures", nargs="*", help="numbers contained in the image file names which are selected " \
                                                            "to be inserted in the existing document. Exemple: -p 1 8 4-6 9-10 or " \
                                                            "-p 1,8, 4-6, 9-10 means the images whose name contain the specified numbers will be added or " \
                                                            "or inserted in ascending number order, in this case 1, 4, 5, 6, 8, 9, 10. " \
                                                            "If this parm is omitted, all the pictures in the curreent " \
                                                            "dir are added or inserted. ")
    args = parser.parse_args(argList)

    return args.document, args.insertionPos, args.pictures


def openExistingOrCreateNewWordDoc(documentName):
    '''
    Opens the passed Word doc documentName located in the current dir. If no document with the passed name
    exist in the current dir, a new empty Word document is created.
    by 1.
    :param userDocumentName:

    :return: either existing or brand new document.
    '''
    documentNameWithExt = documentName + WORD_FILE_EXT

    if curDir.isfile(documentNameWithExt):
        return Document(documentNameWithExt)
    else:
        return Document()


def getFilesInDir(directory):
    '''
    Returns the list of file names contained in the passed directory
    :param directory:
    :return:
    '''
    fileList = []

    for fname in os.listdir(directory):
        path = os.path.join(directory, fname)
        if os.path.isdir(path):
            # skip directories
            continue
        else:
            fileList.append(fname)

    return fileList


def determineInsertionPoint(insertionPos, wordDoc):
    '''
    Returns the 'Heading1' paragraph before which to insert the new images. In case the
    passed insertionPoint is 0 or is greater than the position of the last 'Header1'
    paragraph, None is returned, which indicates that the insertion must occur at the end
    of the document.

    :param insertionPos:
    :param wordDoc:
    :return: existing paragraph or None
    '''
    currentHeadingNumber = 1

    for p in wordDoc.paragraphs:
        if p._p.style != HEADING_ONE_STYLE_NAME_ENGLISH and p._p.style != HEADING_ONE_STYLE_NAME_FRENCH:
            continue

        if insertionPos == currentHeadingNumber:
            return p
        else:
            currentHeadingNumber += 1

    return None


def insertImagesBeforeParagraph(paragraph, imgFileLst):
    '''
    Inserts the images whose file name are in the passed imgFileLst before the passed paragraph.
    Returns the number of inserted images.

    :param paragraph: paragraph
    :param imgFileLst:
    :return number of inserted images
    '''
    imgFileLst.sort(key=sortNumberedStringsFunc, reverse=True)
    insertedImgNumber = len(imgFileLst)

    for imageFileName in imgFileLst:
        imageName = imageFileName.split('.')[0]
        paragraph = paragraph.insert_paragraph_before(getBulletParagraphText(imageName), 'List Bullet')
        paragraph = paragraph.insert_paragraph_before('')
        paragraphRun = paragraph.add_run()

        im = Image.open(imageFileName)
        imgWidthPixel, height = im.size
        imgWidthCm = imgWidthPixel / SCREEN_DPI * 2.54
        paragraphRun.add_picture(imageFileName, width=Cm(min(IMG_MAX_WIDTH, imgWidthCm)))

        paragraph = paragraph.insert_paragraph_before(getTitleParagraphText(imageName), 'Heading 1')

    return insertedImgNumber


def getTitleParagraphText(imageName):
    return imageName + '_title'


def getBulletParagraphText(imageName):
    return imageName + '_bullet'


def createOrUpdateWordDocWithImgInDir(commandLineArgs=None):
    '''
    Python utility to add all the images of a directory to a new Word document in order to facilitate
    documentation creation. The images are added in their file name ascending order.

    *** USAGE ***

    In a command window opened on the dir containing the images, after copying the imgToWordDoc.py file
    in it, simply type

    python imgToWordDoc.py

    or

    python imgToWordDoc.py -d <existing Word doc> -i <insertion position>

    This will create a new Word document whose name is the name of the current dir. In case
    the dir already contains a Word documant with the same name, an incremented number is
    appended to the file name !
    :param commandLineArgs: used only for unit testing only
    '''
    if commandLineArgs == None:
        commandLineArgs = sys.argv[1:]

    userDocumentName, userInsertionPos, imageNumbersToAdd = getCommandLineArgs(commandLineArgs)

    curDir = os.getcwd()

    explodedImageNumbersToAdd = None

    if imageNumbersToAdd:
        explodedImageNumbersToAdd = explodeImageNumbersList(imageNumbersToAdd)

    imgFileLst = filterAndSortImageFileNames(curDir, explodedImageNumbersToAdd)
    doc = None
     
    if userDocumentName:
        #user provided a document name. Either open the existing document
        #or create a new empty one
        if WORD_FILE_EXT in userDocumentName:
            targetWordFileNameNoExt = userDocumentName[:-5]
        else:
            targetWordFileNameNoExt = userDocumentName
        doc = openExistingOrCreateNewWordDoc(targetWordFileNameNoExt)
    else:
        #no document name provided, so name the created word file 
        #using the containing dir name
        if os.name == 'posix':
            targetWordFileNameNoExt = curDir.split('/')[-1]
        else:
            targetWordFileNameNoExt = curDir.split('\\')[-1]
        doc = Document()
        
    targetWordFileNameWithExt = determineUniqueFileName(targetWordFileNameNoExt)
    addedImgNumber = 0
    isInsertionMode = False

    if userInsertionPos != None:
        paragraph = determineInsertionPoint(userInsertionPos, doc)

        if paragraph == None:
            addedImgNumber = addImagesAtEndOfDocument(doc, imgFileLst)
        else:
            addedImgNumber = insertImagesBeforeParagraph(paragraph, imgFileLst)
            isInsertionMode = True
    else:
        setDocMargins(doc)
        addedImgNumber = addImagesAtEndOfDocument(doc, imgFileLst)

    doc.save(targetWordFileNameWithExt)

    if userInsertionPos != None:
        if isInsertionMode:
            resultMsg = "Inserted {} image(s) before header {} in document {} and saved the result to {}.".format(addedImgNumber, userInsertionPos, userDocumentName + WORD_FILE_EXT, targetWordFileNameWithExt)
        else:
            if userInsertionPos == 0:
                resultMsg = "Added {} image(s) at end of document {} and saved the result to {}.".format(
                    addedImgNumber, userDocumentName + WORD_FILE_EXT, targetWordFileNameWithExt)
            else:
                resultMsg = "Added {} image(s) at end of document {} and saved the result to {}. Although insertion position {} was provided, no header paragraph was available at this position and the images were added at the end of the document !".format(
                    addedImgNumber, userDocumentName + WORD_FILE_EXT, targetWordFileNameWithExt, userInsertionPos)
    else:
        resultMsg = "{} file created with {} image(s). Manually add auto numbering to the 'Header 1' / 'Titre 1' style !".format(
            targetWordFileNameWithExt, addedImgNumber)

    print(resultMsg)

    return resultMsg


def addImagesAtEndOfDocument(wordDoc, ascSortedImgFileLst):
    '''
    Add the images in ascSortedImgFileLst at the end of the passed word document
    :param  wordDoc: word document, either newly created or existing, normally already containing
                     images.
    :param  ascSortedImgFileLst: image file name list sorted in ascending order of the number
                                 contained in their name.
    :return: number of images added
    '''
    i = 0

    for imageFileName in ascSortedImgFileLst:
        # ajout d'un titre avant l'image
        imageName = imageFileName.split('.')[0]
        wordDoc.add_heading(getTitleParagraphText(imageName), level=1)

        # ajout de l'image. Si l'image est plus large que la largeur maximale, elle est réduite
        im = Image.open(imageFileName)
        imgWidthPixel, height = im.size
        imgWidthCm = imgWidthPixel / SCREEN_DPI * 2.54
        wordDoc.add_picture(imageFileName, width=Cm(min(IMG_MAX_WIDTH, imgWidthCm)))

        # ajout d'un paragraphe bullet points
        paragraph = wordDoc.add_paragraph(getBulletParagraphText(imageName))
        paragraph.style = 'List Bullet'
        i += 1

    return i


def filterAndSortImageFileNames(containingDir, imageNumbersToAdd =  None):
    '''
    Returns an ordered (ascending) list of image 'jpeg' or 'png' file names whose name
    contains a number and which are located in containingDir. The file names are sorted according
    to the number they contains. If one of the image file names does not contain a number, an
    exception is raised !

    If imageNumbersToAdd is provided, only the images file names containing a number which is
    listed in imageNumbersToAdd are returned.

    :param imageNumbersToAdd: list of integers sorted in ascending order
    :param containingDir:
    :return:
    '''
    fileLst = getFilesInDir(containingDir)
    imgFileLst = list(filter(lambda name: ".jpg" in name or ".png" in name, fileLst))

    if imageNumbersToAdd:
        imgFileLst = filterAccordingToNumber(imgFileLst, imageNumbersToAdd)

    imgFileLst.sort(key=sortNumberedStringsFunc)

    return imgFileLst


def sortNumberedStringsFunc(fileName):
    '''
    Sort strings containing a number according to ascending number order.

    Using this function, a list of file names containing 1.jpg, 11.jpg, 2.jpg, name3.png will
    be ordered so: 1.jpg, 2.jpg, name3.png, 11.jpg !
    :param fileName:
    :return: number in img file name as int
    '''
    nb = re.search(r'(\d+).*', fileName)

    if nb == None:
        raise NameError(
            "Invalid img file name encountered: {0}. Img file names must contain a number for them to be inserted in the right order (depends on this number) !".format(
                fileName))

    return int(nb.group(1))


def setDocMargins(doc):
    sections = doc.sections

    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(LATERAL_MARGIN)
        section.right_margin = Cm(LATERAL_MARGIN)


def determineUniqueFileName(wordFileName):
    '''Verify if a file with same name exists and increment the name by one in this case.

    Ex: if hello.docx exists, returns hello1, hello2, etc

    :param  wordFileName without extention
    :return wordFileName + incremented number (if wordFileName exists in curr dir) +
            word file extention
    '''
    i = 1
    lookupWordFileName = wordFileName

    while curDir.isfile(lookupWordFileName + WORD_FILE_EXT):
        lookupWordFileName = wordFileName + str(i)
        i += 1

    return lookupWordFileName + WORD_FILE_EXT


def explodeImageNumbersList(imageNumberSpec):
    '''
    Returns the unique sorted integers list representing the numbers specified in imageNumberSpec.

    :param imageNumberSpec: list of img nb specs like ['1', '3', '4', '2-7', '9-12'] or, in case of
                            forcing args for argparse in unit testing context ['1 3 4 2-7 9-12']

    :return: like [1, 2, 3, 4, 5, 6, 7, 9, 10, 11, 12]
    '''

    # handling the case when unit testing causes argparse to return a list containing one string
    # for the -p arg instead of a list of picture numer specs
    if len(imageNumberSpec) == 1:
        splittedImageNumberSpec = re.split(r" |,", imageNumberSpec[0])
        if len(splittedImageNumberSpec) > 1:
            imageNumberSpec = splittedImageNumberSpec

    numbersSet = set([])

    for item in imageNumberSpec:
        if '-' in item:
            explodedNumerList = explodeNumberSpec(item)
            numbersSet = numbersSet.union(set(explodedNumerList))
        else:
            if item.isdigit():
                numbersSet.add(int(item))

    sortedNumbers = list(numbersSet)
    sortedNumbers.sort()

    return sortedNumbers


def explodeNumberSpec(dashNumberIntervalSpec):
    '''
    Explode a dashed number interval specification like '2-12' or 19-8' in order
    to return an ascending ordered list of the integers representing the extention
    of the passed interval.

    :param dashNumberIntervalSpec: '2-12' or 19-8' or '2-2'

    :return: sorted integers representing the extention of the passed interval
    '''
    match = re.match(r'(\d+)\s*-\s*(\d+)', dashNumberIntervalSpec)

    # convert tuple to list so that map() can be applied to it
    minMaxList = list(match.groups())

    # map() returns an iterator. Build a list on the iterator
    # so that the resulting list can be sorted
    minMaxList = list(map(lambda x : int(x), minMaxList))

    minMaxList.sort()

    return list(range(minMaxList[0], minMaxList[1] + 1))


def filterAccordingToNumber(fileNameWithNumberLst, numberToKeepList):
    filteredFileNameList = []

    for fileName in fileNameWithNumberLst:
        fileNameNumber = re.search(r'(\d+).*', fileName).group(1)
        if int(fileNameNumber) in numberToKeepList:
            filteredFileNameList.append(fileName)

    return filteredFileNameList


if __name__ == '__main__':
    try:
        createOrUpdateWordDocWithImgInDir()
    except NameError as e:
        print(e)
